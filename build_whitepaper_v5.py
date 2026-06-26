from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Whitepaper FastLedger V5 - Trade Operating System.docx"

NAVY = RGBColor(6, 16, 28)
BLUE = RGBColor(2, 114, 192)
SKY = RGBColor(56, 174, 248)
TEAL = RGBColor(20, 150, 145)
GREEN = RGBColor(22, 145, 100)
GOLD = RGBColor(180, 125, 20)
INK = RGBColor(31, 45, 61)
MUTED = RGBColor(92, 111, 128)
LIGHT = "F4F6F9"
PALE_BLUE = "EAF5FC"
PALE_GREEN = "EAF7F1"
PALE_GOLD = "FFF7E3"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13.5, BLUE, 11, 5),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("FASTLEDGER"), 8.5, BLUE, True, font="Aptos Display")
    set_run(p.add_run("  |  Trade Operating System"), 8.5, MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Whitepaper V5  |  Junio 2026  |  "), 8, MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("WHITEPAPER V5"), 10, SKY, True, font="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("FASTLEDGER"), 32, NAVY, True, font="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("El Sistema Operativo del Comercio Exterior"), 18, BLUE, True, font="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(28)
    set_run(
        p.add_run(
            "Plataforma latinoamericana para coordinar operaciones, inteligencia "
            "aduanera, logística, riesgo, proveedores, documentos y financiamiento."
        ),
        12,
        MUTED,
    )

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    labels = [
        ("PRODUCTO", "Trade Operating System"),
        ("MERCADO INICIAL", "Ecuador y Latinoamérica"),
        ("ESTADO", "MVP técnico en evolución"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, labels):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(label + "\n"), 7.5, BLUE, True)
        set_run(p.add_run(value), 10, NAVY, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Equipo"), 9, MUTED, True)
    team = [
        "Erick Hidalgo - Lider",
        "Julian Chicaiza - Secretario",
        "Baruc Lincango - Disenador",
        "Nicolas Ruiz - Investigador",
        "Francisco Sotomayor - Presentador",
    ]
    for member in team:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        set_run(p.add_run(member), 9.5, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(35)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Hackathon 2026 | Quito, Ecuador | ODS 9"), 9, MUTED, True)
    doc.add_page_break()


def add_callout(doc, title, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), 10.5, accent, True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run(text), 9.8, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2300, 3860, 3200])
    headers = ["Nivel", "Capacidades", "Estado al 26 de junio de 2026"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "DCEAF5")
        p = cell.paragraphs[0]
        set_run(p.add_run(text), 9, NAVY, True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Implementado", "Frontend Trade OS; nueva identidad visual basada en LOGO; dashboard Incoterms; FASTY AI dinamico; autenticacion Supabase; API FastAPI en Render; esquema PostgreSQL con RLS; cliente seguro.", "Codigo publicado en GitHub Pages y backend activo en Render."),
        ("Activado / pendiente de secretos", "Persistencia de operaciones, consultas IA, documentos y proveedores mediante Supabase; Gemini 2.5 Flash; Resend.", "La migracion SQL ya esta lista; requiere cargar service_role, Gemini y Resend en Render para que /health marque true."),
        ("Demostrativo", "Digital Twin, telemetria, proveedores de muestra, eventos en vivo, prediccion de costos y verificacion blockchain.", "Interfaz funcional con datos simulados y arquitectura preparada para integraciones reales."),
        ("Arquitectura objetivo", "Hyperledger Fabric, tracking naviero, aseguradoras, certificadoras, pagos y trade finance.", "Roadmap sujeto a alianzas e integraciones."),
    ]
    for level, capability, state in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (level, capability, state)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(text), 8.7, INK, bold=(cell is cells[0]))
    return table


def add_architecture(doc):
    layers = [
        ("EXPERIENCIA", "GitHub Pages / interfaz Trade OS", "FASTY, dashboard Incoterms, marketplace, autenticacion y Digital Twin"),
        ("API SEGURA", "FastAPI en Render", "Autorizacion, operaciones, documentos, health check y orquestacion"),
        ("INTELIGENCIA", "Gemini 2.5 Flash", "Texto, imagen, PDF, audio y respuesta JSON estructurada"),
        ("DATOS", "Supabase Auth + PostgreSQL + RLS", "Usuarios, operaciones, consultas, documentos, eventos y proveedores"),
        ("COMUNICACION", "Resend", "Correos transaccionales y notificaciones"),
        ("TRAZABILIDAD OBJETIVO", "Hyperledger Fabric", "Hashes documentales, identidades y gobierno empresarial"),
    ]
    table = doc.add_table(rows=0, cols=3)
    set_table_geometry(table, [1900, 2860, 4600])
    fills = [PALE_BLUE, "EDF2F7", PALE_GREEN, "F2F4F7", PALE_GOLD, "F0ECF8"]
    for index, (layer, technology, responsibility) in enumerate(layers):
        cells = table.add_row().cells
        for cell in cells:
            set_cell_shading(cell, fills[index])
        p = cells[0].paragraphs[0]
        set_run(p.add_run(layer), 8.5, BLUE, True)
        p = cells[1].paragraphs[0]
        set_run(p.add_run(technology), 9.2, NAVY, True)
        p = cells[2].paragraphs[0]
        set_run(p.add_run(responsibility), 8.8, INK)


def add_recent_updates(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2300, 3960, 3100])
    headers = ["Frente", "Actualizacion realizada", "Impacto operativo"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "DCEAF5")
        set_run(cell.paragraphs[0].add_run(text), 8.7, NAVY, True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Identidad visual", "Rediseño de colores y vista principal tomando LOGO.png como referencia: fondo oscuro, cian electrico, violeta y acentos de marca.", "La plataforma comunica una identidad mas premium, tecnologica y coherente."),
        ("Backend Render", "FastAPI desplegado en Render y conectado desde GitHub Pages mediante FASTLEDGER_BACKEND_URL.", "El frontend ya puede llamar servicios seguros sin exponer secretos."),
        ("Supabase Auth", "Registro, inicio de sesion, reenvio de verificacion y recuperacion de contraseña con mensajes claros.", "Reduce friccion cuando una cuenta existe pero falta confirmar email o recuperar clave."),
        ("Base de datos", "Script SQL 001_trade_os.sql para perfiles, operaciones, consultas IA, documentos, eventos, proveedores y politicas RLS.", "Permite guardar historial y separar datos por usuario."),
        ("FASTY AI", "Flujo dinamico por texto, imagen, PDF y audio; evita asumir China, drones o Guayaquil si los datos reales indican otro caso.", "Mejora precision del analisis y confianza del usuario."),
        ("Incoterms", "Dashboard interactivo con los 11 Incoterms 2020 y tabla de costos que cambia por responsabilidad de comprador/vendedor.", "Ayuda a comparar escenarios y explicar ahorro frente a una gestion tradicional."),
        ("Pedidos pequeños", "Asesor para importacion por libra y pedidos 4x4.", "Cubre un caso popular de comercio transfronterizo para usuarios finales y emprendedores."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(text), 8.1, INK, bold=(cell is cells[0]))


def add_module_matrix(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1800, 2400, 2960, 2200])
    headers = ["Modulo", "Usuario objetivo", "Resultado", "Estado"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "DCEAF5")
        set_run(cell.paragraphs[0].add_run(text), 8.7, NAVY, True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("FASTY Trade AI", "Importador, exportador, asesor", "Expediente con clasificacion, riesgos, costos, ruta y cronograma.", "Frontend + API Render listos; Gemini depende de variable en Render."),
        ("Dashboard Incoterms", "Importador, asesor, comprador", "Relaciona producto, origen, destino e Incoterm para calcular costos y responsabilidades.", "Implementado con Incoterms 2020 y panel de partida arancelaria."),
        ("Digital Twin", "Operaciones y gerencia", "Vista de contenedor, buque, puerto, aduana, bodega y eventos.", "Prototipo visual; requiere tracking real."),
        ("Supplier Marketplace", "Compras y abastecimiento", "Comparacion de reputacion, certificaciones, MOQ y riesgo.", "Modelo de datos y catalogo semilla."),
        ("Risk Score", "Cumplimiento y finanzas", "Indicador 0-100 por pais, producto, proveedor, valor y permisos.", "Motor determinista implementado; analitica avanzada en roadmap."),
        ("Boveda documental", "Aduanas, seguros y auditoria", "Hash SHA-256 de factura, BL y certificados.", "API implementada; Fabric en roadmap."),
        ("Auth y datos", "Usuarios y administracion", "Registro, login, recuperacion, perfiles, operaciones y consultas por usuario.", "Supabase Auth integrado; persistencia depende de variables secretas en Render."),
        ("Trade Finance", "PyME, banco, fintech", "Credito, factoring, confirming y cartas de credito.", "Concepto e integracion futura."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(text), 8.2, INK, bold=(cell is cells[0]))


def add_roadmap(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1500, 2620, 3080, 2160])
    for cell, text in zip(table.rows[0].cells, ["Fase", "Objetivo", "Entregables", "Criterio de salida"]):
        set_cell_shading(cell, "DCEAF5")
        set_run(cell.paragraphs[0].add_run(text), 8.7, NAVY, True)
    rows = [
        ("0-3 meses", "Cerrar MVP seguro", "Variables Render, Supabase service_role, Gemini, Resend, SMTP confiable, perfiles y operaciones reales.", "Usuarios verificados, expedientes persistentes y health check con integraciones activas."),
        ("3-6 meses", "Validar operacion", "OCR robusto, fuentes arancelarias, proveedores verificados, tracking naviero/courier e Incoterms auditables.", "Primeras operaciones piloto auditables con diferencias estimado vs liquidado."),
        ("6-12 meses", "Escalar servicios", "Seguros, pagos, certificaciones, prediccion de costos y panel administrativo.", "Ingresos recurrentes y procesos medidos."),
        ("12-24 meses", "Red empresarial", "Hyperledger Fabric, trade finance y alianzas regionales.", "Consorcio y trazabilidad multiempresa."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_run(cell.paragraphs[0].add_run(text), 8.3, INK, bold=(cell is cells[0]))


def add_sources(doc):
    sources = [
        ("Servicio Nacional de Aduana del Ecuador (SENAE)", "Para importar; documentos de acompañamiento y proceso de DAI.", "https://www.aduana.gob.ec/servicio-al-ciudadano/para-importar/"),
        ("Servicio Nacional de Aduana del Ecuador (SENAE)", "Para exportar; factura, autorizaciones y certificado de origen cuando corresponda.", "https://www.aduana.gob.ec/servicio-al-ciudadano/para-exportar/"),
        ("Servicio Nacional de Aduana del Ecuador (SENAE)", "Regimenes aduaneros de importacion para consumo y exportacion definitiva.", "https://www.aduana.gob.ec/servicios-para-oces/regimenes-aduaneros/"),
        ("Naciones Unidas", "Objetivo de Desarrollo Sostenible 9: industria, innovacion e infraestructura.", "https://sdgs.un.org/goals/goal9"),
        ("Supabase", "Row Level Security para control de acceso por usuario.", "https://supabase.com/docs/guides/database/postgres/row-level-security"),
        ("Supabase", "Configuracion de URL y redirects para Auth.", "https://supabase.com/docs/guides/auth/redirect-urls"),
        ("International Chamber of Commerce", "Incoterms 2020.", "https://iccwbo.org/business-solutions/incoterms-rules/incoterms-2020/"),
        ("Google AI for Developers", "Gemini API y salidas estructuradas.", "https://ai.google.dev/gemini-api/docs/structured-output"),
        ("Render", "Despliegue de aplicaciones FastAPI.", "https://render.com/docs/deploy-fastapi"),
        ("Resend", "API de correo transaccional.", "https://resend.com/docs/api-reference/emails/send-email"),
    ]
    for index, (owner, title, url) in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(f"{index}. {owner}. "), 8.3, INK, True)
        set_run(p.add_run(f"{title} "), 8.3, INK)
        set_run(p.add_run(url), 7.8, BLUE)


doc = Document()
style_document(doc)
add_running_furniture(doc.sections[0])
add_cover(doc)

doc.add_heading("1. Resumen ejecutivo", level=1)
doc.add_paragraph(
    "FastLedger evoluciona de una calculadora de importaciones a un Trade Operating "
    "System: una plataforma que organiza el ciclo completo del comercio exterior "
    "desde la intención de compra o venta hasta la entrega, el cumplimiento, los "
    "documentos y la evaluación financiera."
)
doc.add_paragraph(
    "La version actual ya combina una interfaz publicada en GitHub Pages con un "
    "backend FastAPI activo en Render, autenticacion Supabase desde el navegador, "
    "un esquema PostgreSQL con RLS, recuperacion de contrasena, dashboard de "
    "Incoterms 2020, asesor 4x4 y una identidad visual alineada al logo de marca."
)
doc.add_paragraph(
    "El producto central es FASTY Trade AI, un agente multimodal que recibe texto, "
    "fotografías, facturas PDF o audio y estructura un expediente operativo. El "
    "resultado esperado incluye producto, posible partida arancelaria, permisos, "
    "riesgos, tributos, ruta, cronograma, proveedores candidatos, transportistas, "
    "financiamiento y supuestos pendientes de validación."
)
add_callout(
    doc,
    "Propuesta de valor",
    "Transformar información dispersa y tareas manuales en una operación comercial "
    "coordinada, trazable y accionable, sin reemplazar la validación de autoridades, "
    "agentes de aduana ni especialistas regulados.",
)
doc.add_heading("Estado del proyecto", level=2)
add_status_table(doc)
doc.add_page_break()
doc.add_heading("Actualizaciones implementadas en la app", level=2)
add_recent_updates(doc)

doc.add_heading("2. Problema y oportunidad", level=1)
doc.add_paragraph(
    "Una operación internacional reúne actores, documentos y decisiones que suelen "
    "vivir en correos, hojas de cálculo, chats y portales desconectados. El usuario "
    "debe coordinar clasificación arancelaria, permisos, seguro, flete, documentos, "
    "pagos, aduana, almacenamiento y entrega sin una visión única del riesgo y costo."
)
doc.add_paragraph(
    "SENAE identifica documentos como el documento de transporte, la factura "
    "comercial, el certificado de origen cuando proceda y otros documentos exigidos "
    "por la autoridad o el organismo regulador. También distingue regímenes como "
    "importación para el consumo y exportación definitiva. FastLedger no sustituye "
    "estos procedimientos: los convierte en tareas, controles y evidencias dentro de "
    "un expediente digital."
)
for text in [
    "Fragmentación de información entre comprador, proveedor, agente, naviera, banco y aseguradora.",
    "Dificultad para estimar el costo aterrizado y comparar escenarios antes de comprometer capital.",
    "Riesgo de documentos incompletos, clasificación incorrecta o permisos identificados demasiado tarde.",
    "Poca visibilidad del embarque y ausencia de una trazabilidad documental compartida.",
    "Acceso limitado de PyMEs a proveedores confiables y productos de financiamiento comercial.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("3. Visión de producto: Trade Operating System", level=1)
doc.add_paragraph(
    "FastLedger propone una capa operativa común para importaciones y exportaciones. "
    "La plataforma integra ocho dominios: aduanas, logística, seguros, pagos "
    "internacionales, certificaciones, financiamiento, proveedores y analítica de "
    "riesgo. Cada operación se modela como un expediente con propietario, estados, "
    "eventos, documentos, resultados de IA y controles de acceso."
)
add_module_matrix(doc)

doc.add_heading("4. FASTY Trade AI", level=1)
doc.add_paragraph(
    "FASTY no se plantea como un chatbot genérico. Es un agente de preparación y "
    "orquestación que convierte una solicitud abierta en datos estructurados y una "
    "secuencia de decisiones. La API utiliza Gemini 2.5 Flash y solicita respuestas "
    "JSON validadas por un esquema Pydantic, reduciendo la ambigüedad entre la IA y "
    "el resto del sistema."
)
doc.add_heading("Flujo multimodal", level=2)
for item in [
    "Texto: interpreta producto, volumen, origen, destino, valor e intención comercial.",
    "Imagen: permite analizar visualmente producto, material y uso como insumo de clasificación.",
    "PDF: extrae datos de factura, incluyendo FOB, peso, Incoterm, proveedor y país de origen.",
    "Audio: transforma una instrucción hablada en una solicitud estructurada.",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("Controles de confianza", level=2)
doc.add_paragraph(
    "El prompt del sistema prohíbe presentar leyes, tasas, proveedores o datos en "
    "tiempo real como hechos cuando no exista una fuente verificable. La respuesta "
    "incluye supuestos y un descargo que exige validar clasificación, permisos y "
    "tributos con SENAE, organismos competentes y un agente de aduana."
)
add_callout(
    doc,
    "Principio de diseño",
    "La IA prepara, compara y alerta. Las decisiones regulatorias, financieras y "
    "contractuales permanecen bajo control humano.",
    PALE_GREEN,
    GREEN,
)

doc.add_heading("Dashboard Incoterms y costos aterrizados", level=2)
doc.add_paragraph(
    "La app incorpora un dashboard interactivo que conecta el producto solicitado "
    "con el Incoterm correspondiente. El calculo contempla EXW, FCA, FAS, FOB, "
    "CFR, CIF, CPT, CIP, DAP, DPU y DDP; ajusta flete, seguro, gastos de origen, "
    "gastos locales, tributos y responsabilidades segun comprador o vendedor."
)
doc.add_paragraph(
    "El panel tambien abre una lectura de partida arancelaria bajo el Sistema "
    "Armonizado, muestra supuestos y explica el ahorro frente a una agencia "
    "tradicional: menos retrabajo documental, comparacion temprana de escenarios, "
    "separacion de costos incluidos por el vendedor y trazabilidad del expediente."
)

doc.add_heading("5. Arquitectura tecnológica", level=1)
doc.add_paragraph(
    "La arquitectura separa la interfaz pública de las credenciales y procesos "
    "sensibles. GitHub Pages entrega el frontend; FastAPI ejecuta la lógica segura; "
    "Supabase gestiona identidad y datos; Gemini procesa el razonamiento multimodal; "
    "Resend envía notificaciones. Las claves secretas existen únicamente como "
    "variables de entorno en Render."
)
add_architecture(doc)
doc.add_heading("Contratos principales del API", level=2)
for item in [
    "GET /health: estado e integraciones configuradas.",
    "GET /v1/me: identidad validada contra Supabase Auth.",
    "POST /v1/operations: creación de una operación propiedad del usuario.",
    "POST /v1/fasty/analyze: análisis estructurado de texto.",
    "POST /v1/fasty/analyze-file: análisis multimodal de imagen, PDF o audio.",
    "POST /v1/documents: registro y hash SHA-256 de documentos.",
    "GET /v1/suppliers: catálogo público de proveedores verificados.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6. Datos, identidad y seguridad", level=1)
doc.add_paragraph(
    "Supabase Auth verifica la sesión del usuario. FastAPI consulta el endpoint de "
    "usuario de Supabase antes de aceptar una operación privada. La base PostgreSQL "
    "incluye perfiles, operaciones, consultas IA, documentos, eventos y proveedores."
)
doc.add_paragraph(
    "Row Level Security agrega defensa en profundidad: cada usuario solo puede leer "
    "o modificar sus propias filas. El backend, aunque utilice una clave administrativa "
    "para tareas de servidor, añade filtros explícitos por user_id a las operaciones "
    "de lectura, actualización y eliminación."
)
doc.add_heading("Autenticacion y persistencia", level=2)
doc.add_paragraph(
    "El flujo de autenticacion fue actualizado para manejar registro existente, "
    "cuentas pendientes de confirmacion, reenvio de verificacion y recuperacion de "
    "contrasena. El objetivo es evitar que el usuario interprete un error generico "
    "de Supabase como si la cuenta no existiera cuando en realidad falta confirmar "
    "el correo o recuperar la clave."
)
doc.add_paragraph(
    "Una vez cargadas las variables de Render, la API puede registrar operaciones, "
    "consultas de FASTY, documentos y eventos en Supabase. El endpoint /health "
    "separa supabase_auth, supabase_database, gemini y resend para diagnosticar "
    "rapidamente que integracion falta activar."
)

doc.add_heading("Protecciones implementadas", level=2)
for item in [
    "Tokens Supabase en cabecera Bearer; secretos fuera del navegador.",
    "CORS restringido al dominio de GitHub Pages y entornos locales autorizados.",
    "Límites de tamaño y tipos MIME permitidos para archivos multimodales.",
    "Hashes SHA-256 para detectar alteraciones documentales.",
    "Validación de modelos de entrada y salida con Pydantic.",
    "Pruebas automatizadas de salud, riesgo y fallback de Gemini.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("7. Digital Twin y trazabilidad blockchain", level=1)
doc.add_paragraph(
    "El Digital Twin representa una operación mediante origen, destino, embarcación, "
    "contenedor, ETA, eventos, documentos y alertas. La interfaz ya permite cambiar "
    "operaciones y capas visuales; la telemetría actual es demostrativa. La fase "
    "productiva requiere integración con navieras, puertos, couriers o agregadores."
)
doc.add_paragraph(
    "La trazabilidad comienza con hashes SHA-256 almacenados junto al expediente. "
    "La arquitectura objetivo utiliza Hyperledger Fabric porque su modelo permissioned "
    "permite identidades conocidas, políticas de acceso, canales y datos privados. "
    "La blockchain registraría huellas y eventos verificables; los archivos sensibles "
    "permanecerían fuera de la cadena bajo controles de acceso."
)
add_callout(
    doc,
    "Documentos priorizados",
    "Factura comercial, Bill of Lading, certificado de origen, certificados "
    "sanitarios o técnicos y póliza de seguro.",
    PALE_GOLD,
    GOLD,
)

doc.add_heading("8. Risk Score y analítica predictiva", level=1)
doc.add_paragraph(
    "Cada operación recibe un indicador de 0 a 100. El motor implementado combina "
    "producto, país, valor FOB y número de permisos como una primera capa "
    "determinista. No pretende ser un modelo de prevención de lavado ni una decisión "
    "crediticia; es un mecanismo de priorización y explicación."
)
doc.add_heading("Evolución prevista", level=2)
for item in [
    "Historial verificado de proveedor, comprador y transportista.",
    "Alertas de sobrevaloración o subvaloración contra referencias autorizadas.",
    "Cambios regulatorios y requisitos por partida.",
    "Predicción de fletes, tipo de cambio, congestión y tiempo de despacho.",
    "Monitoreo de drift, falsos positivos y decisiones humanas posteriores.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("9. Marketplace y financiamiento", level=1)
doc.add_paragraph(
    "El marketplace organiza proveedores por país, categoría, certificaciones, "
    "pedido mínimo, lead time, reputación y riesgo. Los registros actuales son datos "
    "semilla para demostrar el modelo; la etiqueta 'verificado' solo deberá utilizarse "
    "en producción después de un proceso documentado de KYC empresarial, revisión de "
    "certificados, referencias y capacidad productiva."
)
doc.add_paragraph(
    "La capa de Trade Finance se proyecta como un marketplace regulado de opciones: "
    "crédito de importación, factoring, confirming y cartas de crédito. FastLedger "
    "prepararía el expediente y el score operativo; la aprobación corresponde a "
    "bancos, cooperativas o fintech autorizadas."
)

doc.add_page_break()
doc.add_heading("10. Modelo de negocio", level=1)
doc.add_paragraph(
    "El modelo deja de depender de una tarifa única por libra. FastLedger puede "
    "capturar valor mediante software, servicios transaccionales e integraciones, "
    "manteniendo separados los costos logísticos y tributos que dependen de cada caso."
)
table = doc.add_table(rows=1, cols=3)
set_table_geometry(table, [2500, 4260, 2600])
for cell, text in zip(table.rows[0].cells, ["Línea", "Propuesta", "Unidad de ingreso"]):
    set_cell_shading(cell, "DCEAF5")
    set_run(cell.paragraphs[0].add_run(text), 8.8, NAVY, True)
for row in [
    ("SaaS", "Panel, expedientes, colaboración y analítica.", "Suscripción mensual."),
    ("FASTY", "Análisis multimodal y generación de planes.", "Créditos o consumo."),
    ("Operaciones", "Coordinación con socios logísticos y aduaneros.", "Tarifa por expediente."),
    ("Marketplace", "Conexión y verificación de proveedores.", "Comisión o membresía."),
    ("Integraciones", "APIs para empresas, agentes, bancos y aseguradoras.", "Licencia B2B."),
]:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        set_run(cell.paragraphs[0].add_run(text), 8.5, INK, bold=(cell is cells[0]))

doc.add_heading("11. Roadmap de ejecución", level=1)
add_roadmap(doc)
doc.add_heading("Métricas de validación", level=2)
for item in [
    "Porcentaje de expedientes que llegan completos a revisión humana.",
    "Tiempo desde la solicitud hasta un plan operativo utilizable.",
    "Diferencia entre costo estimado y costo liquidado.",
    "Alertas regulatorias confirmadas y descartadas.",
    "Retención de usuarios y operaciones repetidas.",
    "Proveedores verificados con transacciones satisfactorias.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("12. Impacto y ODS 9", level=1)
doc.add_paragraph(
    "El ODS 9 articula infraestructura, industrialización e innovación. FastLedger "
    "contribuye mediante infraestructura digital para PyMEs, estandarización de "
    "expedientes, mejor acceso a información y capacidad de integrar actores del "
    "comercio exterior. El impacto deberá medirse con evidencia operacional y no con "
    "porcentajes aspiracionales sin validación."
)
doc.add_page_break()
doc.add_heading("13. Riesgos del proyecto", level=1)
for item in [
    "Regulación cambiante o fuentes arancelarias no actualizadas.",
    "Alucinaciones o clasificación incorrecta de la IA.",
    "Dependencia de APIs externas y disponibilidad de integraciones.",
    "Protección de datos comerciales, personales y financieros.",
    "Uso indebido de etiquetas de proveedor verificado.",
    "Adopción limitada si no se integra con los flujos reales de agentes y operadores.",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "La mitigación combina fuentes oficiales, trazabilidad de respuestas, revisión "
    "humana, permisos mínimos, RLS, auditoría, contratos con proveedores y despliegues "
    "graduales con usuarios piloto."
)

doc.add_heading("14. Equipo y aprendizaje", level=1)
doc.add_paragraph(
    "La evolución hacia Trade OS obligó al equipo a pasar de una demostración centrada "
    "en cálculo a una arquitectura por capas. El principal aprendizaje es que la "
    "innovación no consiste en añadir IA o blockchain como etiquetas, sino en asignar "
    "a cada tecnología una responsabilidad verificable dentro del proceso."
)
for member in [
    "Erick Hidalgo - liderazgo y dirección del proyecto.",
    "Julian Chicaiza - coordinación y documentación.",
    "Baruc Lincango - diseño de experiencia y comunicación visual.",
    "Nicolas Ruiz - investigación y validación.",
    "Francisco Sotomayor - presentación y vinculación.",
]:
    doc.add_paragraph(member, style="List Bullet")

doc.add_heading("15. Conclusión", level=1)
doc.add_paragraph(
    "FastLedger ya dispone de una experiencia Trade OS, un backend FastAPI probado, "
    "modelos de datos con RLS, autenticación, cliente seguro y contratos para Gemini, "
    "Resend y Render. El siguiente hito no es agregar más pantallas: es activar las "
    "integraciones, ejecutar pilotos reales y medir precisión, tiempos, costo y riesgo."
)
add_callout(
    doc,
    "Tesis",
    "FASTLEDGER puede convertirse en la capa operativa del comercio exterior para "
    "Latinoamérica si combina inteligencia asistida, datos confiables, socios "
    "regulados y trazabilidad auditable.",
    PALE_GREEN,
    GREEN,
)

doc.add_heading("Referencias", level=1)
add_sources(doc)

doc.core_properties.title = "FastLedger V5 - Trade Operating System"
doc.core_properties.subject = "Whitepaper tecnico y de negocio"
doc.core_properties.author = "Equipo FastLedger"
doc.core_properties.keywords = "trade operating system, comercio exterior, inteligencia artificial, FastAPI, Supabase, Gemini, Hyperledger Fabric"

doc.save(OUT)
print(OUT)
